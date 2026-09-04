from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from app.services.text_extractor import TextExtractionError, extract_docx_text, extract_pdf_text


class PdfConversionError(Exception):
    """Raised when a PDF conversion cannot be completed."""


def pdf_to_images_zip(path: Path) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise PdfConversionError("PDF 라이브러리 PyMuPDF가 설치되어 있지 않습니다.") from exc

    try:
        archive = BytesIO()
        with fitz.open(path) as document:
            if document.page_count == 0:
                raise PdfConversionError("비어 있는 PDF는 변환할 수 없습니다.")
            with ZipFile(archive, "w", ZIP_DEFLATED) as zip_file:
                for index, page in enumerate(document, start=1):
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    zip_file.writestr(f"page-{index}.png", pixmap.tobytes("png"))
        return archive.getvalue()
    except PdfConversionError:
        raise
    except Exception as exc:
        raise PdfConversionError("PDF 이미지 변환에 실패했습니다.") from exc


def merge_pdfs(paths: list[Path]) -> bytes:
    if len(paths) < 2:
        raise PdfConversionError("PDF 병합에는 파일이 2개 이상 필요합니다.")

    try:
        import fitz
    except ImportError as exc:
        raise PdfConversionError("PDF 라이브러리 PyMuPDF가 설치되어 있지 않습니다.") from exc

    output = fitz.open()
    try:
        for path in paths:
            with fitz.open(path) as document:
                output.insert_pdf(document)

        if output.page_count == 0:
            raise PdfConversionError("병합할 PDF 페이지가 없습니다.")

        return output.tobytes(garbage=4, deflate=True)
    except PdfConversionError:
        raise
    except Exception as exc:
        raise PdfConversionError("PDF 병합에 실패했습니다.") from exc
    finally:
        output.close()


def split_pdf(path: Path, start_page: int, end_page: int) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise PdfConversionError("PDF 라이브러리 PyMuPDF가 설치되어 있지 않습니다.") from exc

    if start_page < 1 or end_page < start_page:
        raise PdfConversionError("페이지 범위가 올바르지 않습니다.")

    output = fitz.open()
    try:
        with fitz.open(path) as document:
            if end_page > document.page_count:
                raise PdfConversionError(
                    f"마지막 페이지는 {document.page_count} 이하로 지정해야 합니다."
                )
            output.insert_pdf(document, from_page=start_page - 1, to_page=end_page - 1)
        return output.tobytes(garbage=4, deflate=True)
    except PdfConversionError:
        raise
    except Exception as exc:
        raise PdfConversionError("PDF 분할에 실패했습니다.") from exc
    finally:
        output.close()


def compress_pdf(path: Path) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise PdfConversionError("PDF 라이브러리 PyMuPDF가 설치되어 있지 않습니다.") from exc

    try:
        with fitz.open(path) as document:
            if document.page_count == 0:
                raise PdfConversionError("비어 있는 PDF는 압축할 수 없습니다.")
            return document.tobytes(garbage=4, deflate=True, clean=True)
    except PdfConversionError:
        raise
    except Exception as exc:
        raise PdfConversionError("PDF 압축에 실패했습니다.") from exc


def pdf_to_docx(path: Path) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise PdfConversionError("DOCX 라이브러리 python-docx가 설치되어 있지 않습니다.") from exc

    try:
        text = extract_pdf_text(path)
        document = Document()
        document.add_heading("PDF 텍스트 변환 결과", level=1)
        for paragraph in text.splitlines():
            document.add_paragraph(paragraph)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    except Exception as exc:
        raise PdfConversionError("PDF를 DOCX로 변환하지 못했습니다.") from exc


def pdf_to_xlsx(path: Path) -> bytes:
    try:
        import fitz
        from openpyxl import Workbook
    except ImportError as exc:
        raise PdfConversionError("XLSX 변환 라이브러리가 설치되어 있지 않습니다.") from exc

    try:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "PDF Text"
        sheet.append(["page", "line", "text"])

        with fitz.open(path) as document:
            for page_index, page in enumerate(document, start=1):
                text = page.get_text()
                for line_index, line in enumerate(text.splitlines(), start=1):
                    sheet.append([page_index, line_index, line])

        buffer = BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()
    except Exception as exc:
        raise PdfConversionError("PDF를 XLSX로 변환하지 못했습니다.") from exc


def file_to_pdf(path: Path, extension: str) -> bytes:
    if extension in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return text_to_pdf(text, "텍스트 PDF 변환 결과")
    if extension == ".docx":
        try:
            text = extract_docx_text(path)
        except TextExtractionError as exc:
            raise PdfConversionError(str(exc)) from exc
        return text_to_pdf(text, "DOCX PDF 변환 결과")
    if extension in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        return image_to_pdf(path)
    raise PdfConversionError(f"PDF로 변환할 수 없는 파일 형식입니다: {extension}")


def text_to_pdf(text: str, title: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise PdfConversionError("PDF 생성 라이브러리 reportlab이 설치되어 있지 않습니다.") from exc

    try:
        font_name = "HYSMyeongJo-Medium"
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        left = 20 * mm
        top = height - 22 * mm
        line_height = 7 * mm
        max_width = width - (40 * mm)

        pdf.setTitle(title)
        pdf.setFont(font_name, 14)
        pdf.drawString(left, top, title)
        y = top - (12 * mm)
        pdf.setFont(font_name, 10)

        lines = text.splitlines() or [""]
        for line in lines:
            for wrapped in wrap_text(line, font_name, 10, max_width, pdfmetrics):
                if y < 20 * mm:
                    pdf.showPage()
                    pdf.setFont(font_name, 10)
                    y = top
                pdf.drawString(left, y, wrapped)
                y -= line_height

        pdf.save()
        return buffer.getvalue()
    except Exception as exc:
        raise PdfConversionError("텍스트를 PDF로 변환하지 못했습니다.") from exc


def wrap_text(text: str, font_name: str, font_size: int, max_width: float, pdfmetrics) -> list[str]:
    if not text:
        return [""]

    lines: list[str] = []
    current = ""

    for char in text:
        candidate = current + char
        if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = char

    if current:
        lines.append(current)
    return lines


def image_to_pdf(path: Path) -> bytes:
    try:
        from PIL import Image
    except ImportError as exc:
        raise PdfConversionError("이미지 PDF 변환 라이브러리 Pillow가 설치되어 있지 않습니다.") from exc

    try:
        with Image.open(path) as image:
            frames = []
            for frame_index in range(getattr(image, "n_frames", 1)):
                image.seek(frame_index)
                frame = image.convert("RGB")
                frames.append(frame.copy())

        buffer = BytesIO()
        first, *rest = frames
        first.save(buffer, format="PDF", save_all=True, append_images=rest)
        return buffer.getvalue()
    except Exception as exc:
        raise PdfConversionError("이미지를 PDF로 변환하지 못했습니다.") from exc
